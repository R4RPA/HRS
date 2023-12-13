import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from scipy.stats import f_oneway
from docx import Document
from docx.shared import Inches
import io
import os
import env_params as env
from logger_master import logger, log_function_entry_exit


@log_function_entry_exit(logger)
class DataAnalyzer:
    def __init__(self):
        self.env = env
        self.df = None
        self.document = Document()
        self.buffer = io.BytesIO()

    def load_data(self):
        cleaned_file = os.path.join('content', f'cleaned_data_{self.env.start_year}_{self.env.end_year}.csv')
        self.df = pd.read_csv(cleaned_file, index_col=False)
        return self.df.head()

    def add_plot_to_document(self, title, description):
        plt.savefig(self.buffer, format='png')
        self.buffer.seek(0)
        self.document.add_heading(title, level=2)
        self.document.add_paragraph(description)
        self.document.add_picture(self.buffer, width=Inches(6))
        self.buffer = io.BytesIO()
        plt.close()

    def generate_plots(self, x_col, y_col, plot_type, title, description):
        if plot_type == 'boxplot':
            sns.boxplot(data=self.df, x=x_col, y=y_col)
        elif plot_type == 'histplot':
            sns.histplot(self.df[y_col])
        self.add_plot_to_document(title, description)

    def calculate_group_averages(self, group_col, score_cols):
        for score_col in score_cols:
            self.generate_group_average_plot(group_col, score_col)

    def generate_group_average_plot(self, group_col, score_col):
        average_scores = self.df.groupby(group_col)[score_col].mean().reset_index(name='average_score')
        score_counts = self.df[group_col].value_counts().reset_index(name='sample_count')
        result_df = pd.merge(average_scores, score_counts, left_on=group_col, right_on='index', how='left').drop('index', axis=1)
        plt.plot(result_df[group_col], result_df['average_score'], marker='o', label='Average Score')
        plt.xlabel(group_col)
        plt.ylabel(f'Average {score_col}')
        plt.title(f'Average {score_col} vs. {group_col}')
        plt.legend()
        for i, txt in enumerate(result_df['sample_count']):
            plt.annotate(f'{txt}', (result_df[group_col].iloc[i], result_df['average_score'].iloc[i]), fontsize=8)
        title = f'Group Average Plot: {score_col} vs. {group_col}'
        description = f'This line plot shows the average {score_col} for different groups in {group_col}. Sample counts are annotated on the plot.'
        self.add_plot_to_document(title, description)

    def perform_anova(self, group_col, score_col):
        df_dropped = self.df[[group_col, score_col]].dropna()
        groups = df_dropped[group_col].unique()

        # Call f_oneway without explicit type hinting
        anova_results = f_oneway(*(df_dropped[df_dropped[group_col] == group][score_col] for group in groups))

        self.document.add_heading('ANOVA Results', level=2)
        self.document.add_paragraph(f"Analysis of Variance (ANOVA) was performed to compare the means of {score_col} across different groups in {group_col}.")
        self.document.add_paragraph(f"Results: Statistic = {anova_results.statistic}, P-value = {anova_results.pvalue}")

        if anova_results.pvalue < 0.05:
            self.document.add_paragraph("Conclusion: There are significant differences between groups.")
        else:
            self.document.add_paragraph("Conclusion: There are no significant differences between groups.")

    def save_document(self):
        results_doc_file = os.path.join('content', f'analysis_results_{self.env.start_year}_{self.env.end_year}.docx')
        self.document.save(results_doc_file)


#Example Usage:
analyzer = DataAnalyzer()
analyzer.load_data()
analyzer.generate_plots('race_ethnicity', 'cognition_score', 'boxplot', 'Race vs Cognition Score', 'This boxplot shows the distribution of cognition scores across different races.')
analyzer.calculate_group_averages('cognition_score', ['physical_disorder_score', 'final_discrimination_score'])
analyzer.perform_anova('race_ethnicity', 'cognition_score')
analyzer.save_document()
